"""
简历解析服务
支持 PDF、DOCX、图片格式
"""
import os
import re
import zipfile
from typing import Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from PIL import Image
import pytesseract
from loguru import logger


class ResumeParser:
    """简历解析器"""
    
    @staticmethod
    async def parse_file(file_path: str) -> Dict[str, Any]:
        """
        解析简历文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            Dict: {
                "success": bool,
                "text": str,  # 提取的文本
                "error": str,  # 错误信息（如果有）
            }
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == ".pdf":
                text = await ResumeParser._parse_pdf(file_path)
            elif file_ext in [".docx", ".doc"]:
                text = await ResumeParser._parse_docx(file_path)
            elif file_ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
                text = await ResumeParser._parse_image(file_path)
            else:
                return {
                    "success": False,
                    "text": "",
                    "error": f"不支持的文件格式: {file_ext}",
                }
            
            return {
                "success": True,
                "text": text,
                "error": None,
            }
        
        except Exception as e:
            logger.error(f"简历解析失败 ({file_path}): {str(e)}")
            return {
                "success": False,
                "text": "",
                "error": str(e),
            }
    
    @staticmethod
    async def _parse_pdf(file_path: str) -> str:
        """解析 PDF 文件"""
        text_parts = []
        
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                if text:
                    text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"PDF 解析成功 - 文件: {file_path}, 页数: {len(text_parts)}, 字符数: {len(full_text)}")
        
        return full_text
    
    @staticmethod
    async def _parse_docx(file_path: str) -> str:
        """解析 DOCX 文件，支持标准段落、表格、文本框等多种格式"""
        text_parts = []
        
        # 方法1: 使用python-docx解析标准段落和表格
        try:
            doc = docx.Document(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 解析表格
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
        except Exception as e:
            logger.warning(f"python-docx解析失败: {e}")
        
        # 方法2: 如果标准解析结果太少，使用底层XML解析（处理文本框、形状等）
        if len("".join(text_parts)) < 100:
            logger.info(f"标准解析结果较少({len(''.join(text_parts))}字符)，尝试XML底层解析")
            xml_text = await ResumeParser._parse_docx_xml(file_path)
            if len(xml_text) > len("".join(text_parts)):
                text_parts = [xml_text]
                logger.info(f"使用XML底层解析结果: {len(xml_text)}字符")
        
        full_text = "\n".join(text_parts)
        logger.info(f"DOCX 解析成功 - 文件: {file_path}, 字符数: {len(full_text)}")
        
        return full_text
    
    @staticmethod
    async def _parse_docx_xml(file_path: str) -> str:
        """
        通过直接解析docx内部XML来提取文本
        可以处理文本框(textbox)、形状(shape)等非标准段落内容
        """
        all_texts = []
        
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                # 解析主文档
                if 'word/document.xml' in z.namelist():
                    content = z.read('word/document.xml').decode('utf-8')
                    # 提取所有<w:t>标签内的文本（Word文本节点）
                    texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', content)
                    all_texts.extend(texts)
                
                # 解析页眉页脚（可能包含姓名等信息）
                for name in z.namelist():
                    if name.startswith('word/header') or name.startswith('word/footer'):
                        try:
                            content = z.read(name).decode('utf-8')
                            texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', content)
                            all_texts.extend(texts)
                        except:
                            pass
        except Exception as e:
            logger.error(f"XML底层解析失败: {e}")
            return ""
        
        # 合并文本，去除重复
        seen = set()
        unique_texts = []
        for t in all_texts:
            t = t.strip()
            if t and t not in seen:
                seen.add(t)
                unique_texts.append(t)
        
        return " ".join(unique_texts)
    
    @staticmethod
    async def _parse_image(file_path: str) -> str:
        """解析图片文件（OCR）"""
        image = Image.open(file_path)
        
        # 使用 Tesseract OCR
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        
        logger.info(f"图片 OCR 解析成功 - 文件: {file_path}, 字符数: {len(text)}")
        
        return text


# 创建全局实例
resume_parser = ResumeParser()
